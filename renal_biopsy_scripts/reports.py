"""
Functions for interfacing with Pathology Reports
"""

import sys
import os
import json
from pathlib import Path
from typing import Union, List, Dict, Any, OrderedDict, Tuple
import random
from docx import Document

# from simplify_docx import simplify
from torch.utils.data import Dataset

__all__ = ["Report", "ReportsDataset"]


class Report(OrderedDict):
    """ """

    def __init__(self, category: str, report_id: str, filepath: str):
        super().__init__()
        self.category = category
        self.report_id = report_id
        self.filepath = filepath


class ReportsDataset(Dataset):
    """
    Reports Dataset.

    Performs all document parsing.

    Parameters
    ----------
    path_to_reports_dir: Union[Path, str]
        Path the reports directory (parent dir of TCMR and Others)
    init_parser_server: bool
        Initialize the parser server.
        Default: True
    """

    BLACKLIST = ["16SP-21868.docx"]

    def __init__(self, path_to_reports_dir: Union[Path, str]):
        if not self.__is_valid_root(path_to_reports_dir):
            raise ValueError(
                f"{path_to_reports_dir} is not a valid root directory for reports."
            )

        self.root = Path(path_to_reports_dir)
        self.tcmr_reports = [
            i
            for i in os.listdir(self.root / "TCMR")
            if "_" not in i and i not in self.BLACKLIST
        ]
        self.other_reports = [
            i
            for i in os.listdir(self.root / "Others")
            if "_" not in i and i not in self.BLACKLIST
        ]

    def __is_valid_root(self, path_to_reports_dir) -> bool:
        subfolders = os.listdir(path_to_reports_dir)
        return "TCMR" in subfolders and "Others" in subfolders

    def __len__(self):
        return len(self.tcmr_reports) + len(self.other_reports)

    def __getitem__(self, index: int):
        if index < 0 or index > len(self):
            raise IndexError("Index out of range")
        if index < len(self.tcmr_reports):
            report = self.__parse("TCMR", self.tcmr_reports[index])
        else:
            doc = self.other_reports[index - len(self.tcmr_reports)]
            report = self.__parse("Others", doc)
        return report

    def __parse(self, category, filepath: str) -> Report:
        """
        This function is very ad-hoc because it is difficult to parse
        deterministically because the formatting of the tables is
        inconsistent.
        """
        path = self.root / category / filepath
        print("Path: ", path)
        doc = Document(str(path))

        def iter_tables(tables):
            while len(tables) > 0:
                table, tables = tables[0], tables[1:]
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs  # ["\n".join([p.text for p in cell.paragraphs])]
                        yield from iter_tables(cell.tables)
            return None

        i = 0
        d = Report(category, filepath.split(".")[0], filepath)
        signed_and_authorized = False
        tabularized_diagnosis = False
        for p in iter_tables(doc.tables):
            ptext = p.text
            if len(ptext.strip()) == 0:
                continue
            ptext = ptext.strip()
            ptext = ptext.replace("\xa0", " ")
            if "Electronically Signed By" in ptext:
                if "/Authorizing" in ptext:
                    signed_and_authorized = True
            elif "/" in ptext and ":" in ptext and "Date and Time" not in d:
                d["Date and Time"] = ptext
                i += 1
            elif i == 1:
                d["Department"] = ptext
                i += 1
            elif i == 2:
                if signed_and_authorized:
                    d["Electronically Signed By"] = ptext
                    d["Authorizing"] = ptext
                    i = 0
                else:
                    d["Electronically Signed By"] = ptext
                    i += 1
            elif i == 3:
                d["Authorizing"] = ptext
                i = 0
            elif "Order: " in ptext:
                d["Order"] = ptext
            elif "Status: " in ptext:
                d["Status Line"] = ptext
            # See TCMR/19SP-01938.docx for examples on what to do if diagnosis is not inside a single cell
            elif "final diagnosis" in ptext.lower():
                tabularized_diagnosis = True
                d["Diagnosis"] = []
            elif "DIAGNOSIS" in ptext and not tabularized_diagnosis:
                d["Diagnosis"] = ptext
            elif tabularized_diagnosis and "Diagnosis" in d:
                d["Diagnosis"].append(ptext)

        return d


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(f"Usage: python {sys.argv[0]} <path_to_reports_dir>")
        exit(0)
    root = Path(sys.argv[1])
    ds = ReportsDataset(root)
    reports: List[Report] = [random.choice(ds) for _ in range(20)]
    for report in reports:
        with open(f"{report.report_id}.json", "w") as f:
            json.dump(report, f)
